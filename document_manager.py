import PyPDF2
import tkinter as tk
from tkinter import *
from tkinter import filedialog, messagebox, simpledialog
import docx
import docx2pdf 
import docx2txt
from docx.shared import Pt
from fpdf import FPDF
import re

class PDF_Manager:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF Manager")
        self.mylabel=Label(root,text='Select an option from the following',width=30,bg='tan',fg='Black',font='Century 16')
        self.root.config(bg='tan')
        self.mylabel.pack()
        self.split_button = tk.Button(self.root, text="Split PDF", command=self.split_pdf)
        self.split_button.pack(pady=10)
        self.merge_button = tk.Button(self.root, text="Merge PDFs", command=self.merge_pdf)
        self.merge_button.pack(pady=10)
        self.info_button = tk.Button(self.root, text="Get information", command=self.info_pdf)
        self.info_button.pack(pady=10)
        self.convert_btn = tk.Button(self.root, text="Compress PDF", command=self.compress_pdf)
        self.convert_btn.pack(pady=10)
        self.convert_btn = tk.Button(self.root, text="Convert PDF", command=self.convert_pdf)
        self.convert_btn.pack(pady=10)

    def split_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")]) # To select a single file of PDF type.
        if not file_path:
            return # Returns if the dialogbox is closed without choosing any file.
        pdf_file = open(file_path, "rb") # Opening file in binary mode to read non-text files.
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)
        page_range = simpledialog.askstring("Page number range", f"Total number of pages in the pdf = {num_pages}\nEnter page number range\nFormat : x-y\nx = starting page, y = ending page")        
        pattern = '^\d+-\d+$' # String with two integers with one or more digits separated by a hyphen '-'.
        if not re.match(pattern, page_range): # Using match() method of the regular expressions module - 're' for pattern matching.
            messagebox.showinfo('Error','Invalid format entered!')
        if not page_range:
            return
        page_start, page_end = map(int, page_range.split("-"))
        if page_start <= 0 or page_end > num_pages:
            messagebox.showinfo('Error','Out of page number range!')
        pdf_writer = PyPDF2.PdfWriter()
        for page_num in range(page_start - 1, page_end):
            pdf_writer.add_page(pdf_reader.pages[page_num])
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf") # To save the splitted file with default '.pdf' extension.
        if not save_path:
            return
        with open(save_path, "wb") as pdf_output:
            pdf_writer.write(pdf_output)
        messagebox.showinfo('Success','PDF file has been split successfully.')

    def merge_pdf(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("PDF Files", "*.pdf")]) # To select two or more files to merge into a single PDF.
        if not file_paths:
            return
        pdf_writer = PyPDF2.PdfWriter()
        for file_path in file_paths:
            pdf_file = open(file_path, "rb")
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range (len(pdf_reader.pages)):
                pdf_writer.add_page(pdf_reader.pages[page_num])  
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf") # To save the merged file with default '.pdf' extension.
        if not save_path:
            return
        with open(save_path, 'wb') as pdf_output:
            pdf_writer.write(pdf_output)
        messagebox.showinfo("Success", "PDF files have been merged successfully.")    

    def info_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            info = pdf_reader.metadata
            d1 = dict() # Creating a dictionary to store metadata of the pdf file.
            d1['Title'] = info.title
            d1['Author'] = info.author
            d1['Page Count'] = len(pdf_reader.pages)
            d1['Subject'] = info.subject
            d1['Creator'] = info.creator
            d1['Producer'] = info.producer    
        for key, value in d1.items():
            print(f'\n{key} : {value}')

    def compress_pdf(self):
        file_path = filedialog.askopenfilename(filetype=[('PDF Files', '*.pdf')])
        if not file_path:
            return
        pdf_file = open(file_path, 'rb')
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        pdf_writer = PyPDF2.PdfWriter()
        for page in pdf_reader.pages:
            page.compress_content_streams()
            pdf_writer.add_page(page)
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf")
        if not save_path:
            return
        with open(save_path, "wb") as pdf_output:
            pdf_writer.write(pdf_output)
        messagebox.showinfo("Success", 'PDF file has been compressed successfully!')

    def convert_pdf(self): 
        file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        output_format = simpledialog.askstring("Output Format", "Enter output file format (e.g: docx, txt):")
        if not output_format:
            return
        if output_format == "docx":
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                word_document = docx.Document()
                for page_num in range (len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    word_document.add_paragraph(text)
            save_path = filedialog.asksaveasfilename(defaultextension=".docx")   
            word_document.save(save_path)
            messagebox.showinfo("Success", "PDF file has been converted into word document successfully.")        
        elif output_format == 'txt':
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                save_path = filedialog.asksaveasfilename(defaultextension=".txt")   
                text_file = open(save_path, 'w', encoding = 'utf-8')
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_file.write(text)
                text_file.close()
            messagebox.showinfo("Success", "PDF file has been converted into text file successfully.")

class Word_Manager:
    def __init__(self,root):
        self.root = root
        self.root.title('Word Manager')
        self.mylabel=Label(root,text='Select an option from the following',width=30,bg='tan',fg='Black',font='Century 16')
        self.root.config(bg='tan')
        self.mylabel.pack()
        self.split_button = tk.Button(self.root, text="Word to PDF", command=self.word_to_pdf)
        self.split_button.pack(pady=10)
        self.merge_button = tk.Button(self.root, text="Word to Text", command=self.word_to_txt)
        self.merge_button.pack(pady=10)
    def word_to_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if not file_path:
            return
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf")
        if not save_path:
            return      
        docx2pdf.convert(file_path,save_path) # 'convert' function imported from docx2pdf module that takes the paths of word file and pdf file for conversion.
        messagebox.showinfo("Success", "Word file has been converted into PDF successfully.")
    def word_to_txt(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if not file_path:
            return        
        save_path = filedialog.asksaveasfilename(defaultextension=".txt")
        if not save_path:
            return        
        text = docx2txt.process(file_path) # 'process' function imported from 'docx2txt' module returns a string containing the text content of the word file.
        with open(save_path, "w", encoding = 'utf-8') as f:
            f.write(text)
        messagebox.showinfo("Success", "Word file has been converted into text file successfully.")

class Text_Manager():
    def __init__(self, root):
        self.root = root
        self.root.title("Text Manager")
        self.mylabel=Label(root,text='Select an option from the following',width=30,bg='tan',fg='Black',font='Century 16')
        self.root.config(bg='tan')
        self.mylabel.pack()
        self.split_button = tk.Button(self.root, text="Text to PDF", command=self.text_to_pdf)
        self.split_button.pack(pady=10)
        self.merge_button = tk.Button(self.root, text="Text to Word", command=self.text_to_word)
        self.merge_button.pack(pady=10)

    def text_to_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
        if not file_path:
            return
        pdf = FPDF() # Creating a new pdf file.
        pdf.add_page() # Addfing a new page to the pdf.
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf")
        root = tk.Tk()
        root.title('Details')
        myFrame = Frame(root)
        myLabel1 = Label(myFrame, text = 'Font ')
        myLabel2 = Label(myFrame, text = 'Size ')
        myEntry1 = Entry(myFrame)
        myEntry2 = Entry(myFrame)
        myLabel1.grid(row = 0, column = 0)
        myLabel2.grid(row = 1, column = 0)
        myEntry1.grid(row = 0, column = 1)
        myEntry2.grid(row = 1, column = 1)
        def setfont():
            font1 = myEntry1.get()
            if not font1:
                messagebox.showerror("Error", "Invalid font name.")
                return
            size1 = int(myEntry2.get())     
            pdf.set_font(font1, size=size1) # 'set_font' method takes the font (string) and size (integer) as arguments.
            with open(file_path, "r", newline = '\n') as txt_file:
                text = txt_file.readlines()
                for line in text:
                    pdf.multi_cell(w=190,h=8,txt = line)
            pdf.output(save_path)
            messagebox.showinfo("Success", "Text file has been converted into PDF successfully.")
        myButton = Button(myFrame, text = 'Submit', command = setfont)
        myButton.grid(row = 2, columnspan = 2)      
        myFrame.grid(row = 0, column = 0)
        root.mainloop()

    def text_to_word(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
        if not file_path:
            return    
        save_path = filedialog.asksaveasfilename(defaultextension=".docx")
        if not save_path:
            return
        document  = docx.Document() # Creating a new word document.      
        root = tk.Tk()
        root.title('Details')
        myFrame = Frame(root)
        myLabel1 = Label(myFrame, text = 'Font')
        myLabel2 = Label(myFrame, text = 'Size')
        myEntry1 = Entry(myFrame)
        myEntry2 = Entry(myFrame)
        myLabel1.grid(row = 0, column = 0)
        myLabel2.grid(row = 1, column = 0)
        myEntry1.grid(row = 0, column = 1)
        myEntry2.grid(row = 1, column = 1)

        def setfont():  
            style = document.styles['Normal'] # 'styles' attribute to access the styles defined in a document.
            font = style.font
            font.name = myEntry1.get()
            font.size = Pt(int(myEntry2.get())) # 'Pt()' function imported from 'docx.shared' module to convert the font size in pixels to points
            with open(file_path, 'r') as txt_file:
                text = txt_file.read()
            document.add_paragraph(text) # Adding paragraphs to the new word document.
            document.save(save_path)
            messagebox.showinfo("Success", "Text file has been converted into Word document successfully.")             
        myButton = Button(myFrame, text = 'Submit', command = setfont)
        myButton.grid(row = 2, columnspan = 2)      
        myFrame.grid(row = 0, column = 0)        
        root.mainloop()

if __name__ == "__main__":
    n = int(input('1. PDF\n2. Word Document\n3. Text File\n0. Exit\nEnter your choice of file type to access : '))
    if n==1:
        root = tk.Tk()
        pdf_obj = PDF_Manager(root)
        root.mainloop()
    elif n==2:
        root = tk.Tk()
        word_obj = Word_Manager(root)
        root.mainloop()
    elif n==3:
        root = tk.Tk()
        txt_obj = Text_Manager(root)
        root.mainloop()